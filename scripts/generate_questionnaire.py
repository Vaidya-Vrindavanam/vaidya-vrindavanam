from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_title_page(doc):
    """Create professional title page with instructions"""

    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("Vaidya Vrindavanam")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(44, 62, 80)  # Dark blue-gray

    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("About Us Discovery Questionnaire")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(52, 73, 94)

    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Add purpose section
    purpose_heading = doc.add_paragraph()
    purpose_heading_run = purpose_heading.add_run("Purpose")
    purpose_heading_run.font.size = Pt(12)
    purpose_heading_run.font.bold = True

    purpose_text = doc.add_paragraph(
        "This questionnaire is designed to gather comprehensive information about Vaidya Vrindavanam's "
        "mission, philosophy, team, facility, treatments, and patient outcomes. Your detailed responses "
        "will be used to create a compelling About Us page that resonates with both local patients and "
        "international wellness travelers."
    )
    purpose_text.paragraph_format.line_spacing = 1.5

    # Add how to use section
    doc.add_paragraph()
    how_heading = doc.add_paragraph()
    how_heading_run = how_heading.add_run("How to Use This Questionnaire")
    how_heading_run.font.size = Pt(12)
    how_heading_run.font.bold = True

    how_text = doc.add_paragraph(
        "1. You can complete this questionnaire at your own pace — tackle one section at a time if needed.\n"
        "2. Answer each question as fully and honestly as possible.\n"
        "3. Questions marked with [*] are open-ended narrative prompts where we'd love to hear your story.\n"
        "4. Don't worry about perfect grammar — we'll polish the copy once we have your information.\n"
        "5. If you run out of space in any section, continue on the overflow pages at the end (labeled 'Overflow')."
    )
    how_text.paragraph_format.line_spacing = 1.5

    # Add time estimate
    doc.add_paragraph()
    time_para = doc.add_paragraph()
    time_run = time_para.add_run("Estimated Completion Time: ")
    time_run.font.bold = True
    time_para.add_run("2-3 hours (flexible — can be done over multiple sessions)")
    time_para.paragraph_format.line_spacing = 1.5

    # Add page break
    doc.add_page_break()

def add_section_header(doc, section_num, section_title):
    """Add a formatted section header"""
    header = doc.add_paragraph()
    header.paragraph_format.space_before = Pt(12)
    header.paragraph_format.space_after = Pt(6)
    header_run = header.add_run(f"Section {section_num}: {section_title}")
    header_run.font.size = Pt(14)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(44, 62, 80)

def add_question_number(doc, question_num, question_text):
    """Add a numbered question with consistent formatting"""
    q = doc.add_paragraph(style='List Number')
    q.paragraph_format.space_before = Pt(6)
    q.paragraph_format.space_after = Pt(3)
    q_run = q.add_run(question_text)
    q_run.font.size = Pt(11)

def add_structured_response_space(doc, placeholder_text, height_lines=2):
    """Add space for structured responses (short answers, checkboxes)"""
    response = doc.add_paragraph(placeholder_text)
    response.paragraph_format.left_indent = Inches(0.5)
    response.paragraph_format.space_after = Pt(12)
    response_run = response.runs[0]
    response_run.font.italic = True
    response_run.font.color.rgb = RGBColor(149, 165, 166)  # Light gray placeholder
    response_run.font.size = Pt(10)

def add_narrative_prompt(doc, prompt_text):
    """Add space for open-ended narrative response"""
    doc.add_paragraph()
    prompt = doc.add_paragraph()
    prompt_run = prompt.add_run(f"[*] Narrative Prompt: {prompt_text}")
    prompt_run.font.italic = True
    prompt_run.font.bold = True
    prompt_run.font.color.rgb = RGBColor(41, 128, 185)  # Blue for emphasis

    # Add expandable text box instruction
    instruction = doc.add_paragraph("(Please provide a 200-400 word response in the space below)")
    instruction_run = instruction.runs[0]
    instruction_run.font.italic = True
    instruction_run.font.size = Pt(9)
    instruction_run.font.color.rgb = RGBColor(127, 140, 141)

    # Add space for response (gray shaded box simulation)
    response_box = doc.add_paragraph()
    response_box.paragraph_format.left_indent = Inches(0.5)
    for i in range(6):  # 6 lines for text box
        doc.add_paragraph()
    response_box.paragraph_format.space_after = Pt(12)

def add_section_1_origins_mission(doc):
    """Section 1: Origins & Mission"""
    add_section_header(doc, 1, "Origins & Mission")

    doc.add_paragraph(
        "Capture the founding story, 'why we exist,' and current vision for the clinic."
    ).runs[0].font.italic = True

    # Q1
    add_question_number(doc, 1, "How was Vaidya Vrindavanam founded?")
    add_structured_response_space(doc, "Who founded it? When (year)? What inspired the founder/leadership to open this clinic?")

    # Q2
    add_question_number(doc, 2, "How long have you been operating?")
    add_structured_response_space(doc, "Years in operation: _____", height_lines=1)

    # Q3
    add_question_number(doc, 3, "What is your mission statement or core purpose?")
    add_structured_response_space(doc, "(Or: 'What problem do you solve for patients?') — 1-2 sentence mission statement")

    # Q4
    add_question_number(doc, 4, "What are your long-term goals for the clinic?")
    add_structured_response_space(doc, "3-5 bullet points or paragraph format")

    # Narrative Prompt
    add_narrative_prompt(
        doc,
        "Tell us the story of your clinic in 3-4 sentences — what makes it special to you and your patients?"
    )

    doc.add_page_break()

def add_section_2_philosophy_approach(doc):
    """Section 2: Philosophy & Approach"""
    add_section_header(doc, 2, "Philosophy & Approach")

    doc.add_paragraph(
        "Articulate Ayurvedic philosophy and treatment methodology that guides your practice."
    ).runs[0].font.italic = True

    # Q1
    add_question_number(doc, 1, "What is your approach to Ayurveda?")
    doc.add_paragraph("☐ Classical/traditional Ayurveda (as documented in ancient texts)")
    doc.add_paragraph("☐ Integrative (Ayurveda + modern medicine combined)")
    doc.add_paragraph("☐ Modern science-based Ayurveda (evidence-driven)")
    doc.add_paragraph("☐ Other: _________")

    # Q2
    add_question_number(doc, 2, "What are your 3-5 core values?")
    add_structured_response_space(doc, "List core values with brief explanations (e.g., 'Patient-centered care', 'Respect for nature')")

    # Q3
    add_question_number(doc, 3, "How do you personalize treatment for each patient?")
    add_structured_response_space(doc, "Include: assessment process, individualization approach, consultation duration (150-200 words)")

    # Q4
    add_question_number(doc, 4, "Do you integrate any modern medical practices? If so, how?")
    doc.add_paragraph("Yes / No")
    add_structured_response_space(doc, "If yes, explain the integration (imaging, lab work, collaborative approach with allopathy, etc.)")

    # Q5
    add_question_number(doc, 5, "What do you believe makes Ayurvedic treatment effective?")
    add_structured_response_space(doc, "Open-ended narrative (150-250 words)")

    # Narrative Prompt
    add_narrative_prompt(
        doc,
        "Describe your treatment philosophy in plain language — what should patients expect when they come to you?"
    )

    doc.add_page_break()

def add_section_3_team_credentials(doc):
    """Section 3: Team & Credentials"""
    add_section_header(doc, 3, "Team & Credentials")

    doc.add_paragraph(
        "Build credibility and trust through doctor/staff expertise and qualifications."
    ).runs[0].font.italic = True

    # Q1
    add_question_number(doc, 1, "How many doctors/practitioners do you have on staff?")
    add_structured_response_space(doc, "Total count: ___", height_lines=1)

    # Q2
    add_question_number(doc, 2, "For each key practitioner, provide:")
    doc.add_paragraph("Name: ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Qualification(s): ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Years of experience: ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Specialization(s): ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Bio/background (2-3 sentences): ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph()

    # Q3
    add_question_number(doc, 3, "What certifications or training do your practitioners have?")
    add_structured_response_space(doc, "(BAMS, BAMS+specialization, Ayurveda diploma, other certifications, etc.) — Bulleted list")

    # Q4
    add_question_number(doc, 4, "Are there any renowned practitioners, specialists, or notable credentials on your team?")
    add_structured_response_space(doc, "Highlight any notable achievements, publications, recognitions, training abroad, etc.")

    # Q5
    add_question_number(doc, 5, "How does your team approach patient care?")
    doc.add_paragraph("☐ One-on-one consultations with dedicated doctor")
    doc.add_paragraph("☐ Group/team-based assessments")
    doc.add_paragraph("☐ Follow-up adjustments during stay")
    doc.add_paragraph("☐ Collaborative care (multiple specialists)")
    doc.add_paragraph("☐ Other: __________")

    # Narrative Prompt
    add_narrative_prompt(
        doc,
        "What makes your team qualified to treat patients? Share insights on their training, experience, and what patients appreciate about working with them."
    )

    doc.add_page_break()

def add_section_4_facility_experience(doc):
    """Section 4: Facility & Experience"""
    add_section_header(doc, 4, "Facility & Experience")

    doc.add_paragraph(
        "Describe the patient experience — physical space, comfort, amenities, daily life during treatment."
    ).runs[0].font.italic = True

    # Q1
    add_question_number(doc, 1, "Where is your clinic located?")
    doc.add_paragraph("Address: ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("What are the advantages of your location?").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("☐ Proximity to transport/airport").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("☐ Peaceful, natural environment").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("☐ Access to specific resources (water, herbs, climate)").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("☐ Local culture/tourism accessibility").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("☐ Other: __________").paragraph_format.left_indent = Inches(0.75)

    # Q2
    add_question_number(doc, 2, "Facility specifications:")
    doc.add_paragraph("Number of treatment rooms: ___").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Number of patient beds: ___").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Year built/last renovated: ___").paragraph_format.left_indent = Inches(0.5)

    # Q3
    add_question_number(doc, 3, "Do you offer accommodation?")
    doc.add_paragraph("Yes / No").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("If yes:").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Number of rooms: ___").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Room types (single, double, deluxe, etc.): ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Special features (AC, attached bath, garden view, etc.): ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Any special amenities: ___________").paragraph_format.left_indent = Inches(0.75)

    # Q4
    add_question_number(doc, 4, "What does a typical day look like for a patient?")
    doc.add_paragraph("Treatment times: ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Meal times and dietary approach: ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Rest/relaxation periods: ___________").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Evening activities: ___________").paragraph_format.left_indent = Inches(0.5)

    # Q5
    add_question_number(doc, 5, "What amenities do you offer?")
    doc.add_paragraph("☐ Yoga classes")
    doc.add_paragraph("☐ Meditation/pranayama")
    doc.add_paragraph("☐ Garden/outdoor spaces")
    doc.add_paragraph("☐ Spa/wellness facilities")
    doc.add_paragraph("☐ Library/reading areas")
    doc.add_paragraph("☐ Kitchen facilities")
    doc.add_paragraph("☐ WiFi/tech amenities")
    doc.add_paragraph("☐ Other: __________")

    # Narrative Prompt
    add_narrative_prompt(
        doc,
        "What does a typical day look like for a patient staying with you? Paint a picture of their experience from morning to evening."
    )

    doc.add_page_break()

def add_section_5_treatments_specializations(doc):
    """Section 5: Treatments & Specializations"""
    add_section_header(doc, 5, "Treatments & Specializations")

    doc.add_paragraph(
        "Highlight what you do best, unique offerings, and what sets you apart competitively."
    ).runs[0].font.italic = True

    # Q1
    add_question_number(doc, 1, "Which 3-5 treatments are you most known for or specialize in?")
    add_structured_response_space(doc, "List with brief description of each")

    # Q2
    add_question_number(doc, 2, "Which health conditions do you have the most success treating?")
    add_structured_response_space(doc, "List top 5-7 conditions")

    # Q3
    add_question_number(doc, 3, "Are there any rare, unique, or signature treatments you offer that competitors don't?")
    add_structured_response_space(doc, "Narrative section (150-250 words)")

    # Q4
    add_question_number(doc, 4, "Do you offer wellness packages or preventative care programs?")
    doc.add_paragraph("Yes / No").paragraph_format.left_indent = Inches(0.5)
    add_structured_response_space(doc, "If yes, describe the packages and their duration/focus")

    # Q5
    add_question_number(doc, 5, "What makes your treatment approach different from other Ayurvedic clinics?")
    add_structured_response_space(doc, "Key differentiators, unique methodology, special focus areas")

    # Narrative Prompt
    add_narrative_prompt(
        doc,
        "What are you proudest of in your treatment offerings? What do patients most often come to you for, and why do they choose you over other options?"
    )

    doc.add_page_break()

def add_section_6_outcomes_transformations(doc):
    """Section 6: Outcomes & Patient Transformations"""
    add_section_header(doc, 6, "Outcomes & Patient Transformations")

    doc.add_paragraph(
        "Capture success stories and build social proof through patient transformations."
    ).runs[0].font.italic = True

    # Q1
    add_question_number(doc, 1, "Patient Transformation Stories — Please share 2-3 specific examples:")
    doc.add_paragraph("For each:").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Patient name (or anonymous): ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Initial condition/complaint: ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Duration of treatment: ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Results/improvements: ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Current status/long-term outcome: ___________").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Permission to use this story publicly? Yes / No / Anonymous only").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Detailed narrative (200-300 words per story):").paragraph_format.left_indent = Inches(0.75)
    for i in range(7):
        doc.add_paragraph().paragraph_format.left_indent = Inches(1)
    doc.add_paragraph()

    # Q2
    add_question_number(doc, 2, "Do you track success/improvement rates by condition?")
    doc.add_paragraph("Yes / No").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("If yes, share any metrics:").paragraph_format.left_indent = Inches(0.5)
    doc.add_paragraph("Condition: _____ | % improvement: _____").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("Condition: _____ | % improvement: _____").paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph("(Repeat as needed)").paragraph_format.left_indent = Inches(0.75)

    # Q3
    add_question_number(doc, 3, "Do you have patients who return repeatedly or refer others?")
    doc.add_paragraph("Yes / No").paragraph_format.left_indent = Inches(0.5)
    add_structured_response_space(doc, "If yes, what do they say about returning? Why do they refer? (150-250 words)")

    # Q4
    add_question_number(doc, 4, "What do patients typically say about their experience?")
    add_structured_response_space(doc, "Collect any quotes, common feedback themes, or testimonial insights (bulleted list or narrative)")

    # Narrative Prompt
    add_narrative_prompt(
        doc,
        "Tell us about a patient whose life was truly transformed by treatment at your clinic. What was their journey, and what do they say about their experience?"
    )

    doc.add_page_break()

def add_overflow_pages(doc):
    """Add overflow pages for continuation of responses"""
    overflow_header = doc.add_paragraph()
    overflow_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    overflow_run = overflow_header.add_run("OVERFLOW PAGES")
    overflow_run.font.size = Pt(14)
    overflow_run.font.bold = True
    overflow_run.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()

    instructions = doc.add_paragraph(
        "If you need additional space to continue any responses, please indicate which section/question you are continuing "
        "and use the space below. You can continue to as many overflow pages as needed."
    )
    instructions.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()
    continuation_label = doc.add_paragraph()
    continuation_label_run = continuation_label.add_run("Which section/question are you continuing?")
    continuation_label_run.font.bold = True
    continuation_label.paragraph_format.space_before = Pt(12)

    # Add several blank overflow pages
    for page_num in range(1, 4):  # 3 overflow pages
        doc.add_paragraph()
        page_label = doc.add_paragraph(f"Overflow Page {page_num}")
        page_label_run = page_label.runs[0]
        page_label_run.font.italic = True
        page_label_run.font.size = Pt(10)
        page_label_run.font.color.rgb = RGBColor(149, 165, 166)

        # Add writing space
        for i in range(30):
            doc.add_paragraph()

        if page_num < 3:
            doc.add_page_break()

# Initialize document with default settings
doc = Document()

# Call title page function
add_title_page(doc)

# Call section functions
add_section_1_origins_mission(doc)
add_section_2_philosophy_approach(doc)
add_section_3_team_credentials(doc)
add_section_4_facility_experience(doc)
add_section_5_treatments_specializations(doc)
add_section_6_outcomes_transformations(doc)
add_overflow_pages(doc)

# Save the document
output_path = "docs/About-Us-Discovery-Questionnaire.docx"
doc.save(output_path)
print(f"[OK] Questionnaire generated successfully: {output_path}")

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0, 0, 0)

# Set margins (1 inch all around)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

print("Document initialized successfully")
